# Author:ADMIN
# Date:2020/3/22
import sys
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout,
                             QPushButton, QMainWindow)
from mainwindow import Ui_MainWindow

from word_to_excel import WordToExcel
from openpyxl import *
from docx import Document
import re


class MainWindow:
    def __init__(self):

        self.qMainWindow = QMainWindow()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self.qMainWindow)
        self.ui.close.clicked.connect(self.on_button_click)
        self.ui.start.clicked.connect(self.on_button_start)
        self.content = []
        self.pattern_stem = re.compile(r"(\d+\.).*")
        self.pattern_patition = re.compile(r"([A-D](\.|\．){1})(\s?)(\S.+)")



    def show(self):
        self.qMainWindow.show()

    def on_button_click(self):
        sender = self.qMainWindow.sender()
        msg = sender.text() + " 被按下了"
        status = self.qMainWindow.statusBar()
        status.showMessage(msg, 5000)
        # 程序退出
        app = QApplication.instance()
        app.quit()
    def read_ques(self):
        #ques = self.ui.ques.text()
        ques = self.ui.ques.text()
        return ques

    def on_button_start(self):
        str = self.read_ques()
        wordToExcel = WordToExcel()

        stem_file = self.ui.ques.text()
        anwser_file = self.ui.asr.text()
        excel_file = self.ui.dst.text()

        print(type(stem_file))
        stem_doc = Document(stem_file)
        answer_doc = Document(anwser_file).tables
        stem = wordToExcel.get_ques_stem(stem_doc, self.pattern_stem)
        option = wordToExcel.get_option(stem_doc, self.pattern_patition)
        anwser = wordToExcel.get_anwser(answer_doc)

        stem_num = len(stem)
        offset_a = 0
        offset_b = 1
        offset_c = 2
        offset_d = 3
        for num in range(0, stem_num):
            if num < 600:

                tmp_lst = ["选择题",
                           stem[num],
                           option[offset_a],
                           option[offset_b],
                           option[offset_c],
                           option[offset_d],
                           anwser[num]]
            else:
                tmp_lst = ["判断题", stem[num], ' ', ' ', ' ', ' ', anwser[num]]

            offset_a += 4
            offset_b += 4
            offset_c += 4
            offset_d += 4
            print(tmp_lst)
            num += 1
            self.content.append(tmp_lst)
        print('start set excel')
        wordToExcel.set_excel(self.content, stem_num, excel_file)



if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()

    sys.exit(app.exec_())

